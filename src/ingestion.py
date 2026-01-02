import os
import time
import shutil
from typing import List
from PIL import Image
import pytesseract
from pypdf import PdfReader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

# Import our helper functions
from src.utils import setup_logger, compute_md5_hash, is_valid_text

# Create a logger for this file
logger = setup_logger("ingestion")

class DocumentProcessor:
    def __init__(self):
        # We need a place to split our text. 
        # 1000 characters is a good size for context, and 200 overlap helps keeps context between chunks.
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Using HuggingFace because it's free and runs locally.
        # "all-MiniLM-L6-v2" is a popular small model.
        logger.info("Loading embedding model all-MiniLM-L6-v2...")
        self.embeddings = HuggingFaceEmbeddings(
            model_name="all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'}
        )
        
        # We'll stick the database in a local folder called 'chroma_db'
        self.persist_directory = "chroma_db"

    def process_file(self, file_path: str) -> List[Document]:
        """
        Reads a file and returns a list of text chunks (Documents).
        Handles PDF, Images (OCR), and Text files.
        """
        filename = os.path.basename(file_path)
        logger.info(f"Processing file: {filename}")
        
        ext = filename.split('.')[-1].lower()
        raw_text_pages = [] # List of tuples: (page_content, page_number)

        try:
            # 1. READ CONTENT BASED ON TYPE
            if ext == "pdf":
                reader = PdfReader(file_path)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text:
                        raw_text_pages.append((text, i + 1))
                        
            elif ext in ["png", "jpg", "jpeg"]:
                # Use Tesseract for OCR
                try:
                    text = pytesseract.image_to_string(Image.open(file_path))
                    raw_text_pages.append((text, 1)) # Images are treated as page 1
                except Exception as e:
                    logger.error(f"OCR failed for {filename}. Do you have Tesseract installed? Error: {e}")
                    
            elif ext in ["txt", "md"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    text = f.read()
                    raw_text_pages.append((text, 1))

            elif ext == "docx":
                import docx
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                text = "\n".join(full_text)
                if text:
                    raw_text_pages.append((text, 1))
            
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return []

            # 2. CHUNK AND CLEAN
            final_docs = []
            seen_hashes = set() # To keep track of duplicates

            for page_content, page_num in raw_text_pages:
                # Split the page into smaller chunks
                chunks = self.text_splitter.split_text(page_content)
                
                for i, chunk_text in enumerate(chunks):
                    # GUARDRAIL: Skip garbage text
                    if not is_valid_text(chunk_text):
                        continue
                        
                    # GUARDRAIL: Deduplication using MD5
                    chunk_hash = compute_md5_hash(chunk_text)
                    if chunk_hash in seen_hashes:
                        logger.info(f"Skipping duplicate chunk in {filename}")
                        continue
                    seen_hashes.add(chunk_hash)
                    
                    # Create the Document object with required metadata
                    doc = Document(
                        page_content=chunk_text,
                        metadata={
                            "source": filename,
                            "page_number": page_num,
                            "chunk_id": f"{filename}_{page_num}_{i}", # Unique ID for tracing
                            "chunk_hash": chunk_hash
                        }
                    )
                    final_docs.append(doc)

            logger.info(f"Generated {len(final_docs)} chunks from {filename}")
            return final_docs

        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")
            return []

    def create_vector_db(self, documents: List[Document]):
        """
        Takes all documents, clears old DB (if we want a fresh start), 
        and creates a new Chroma vector store.
        """
        if not documents:
            logger.warning("No documents to add to Vector DB.")
            return None

        logger.info("Creating/Updating Vector Database...")
        
        # Simple timer to see how long embedding takes
        start_time = time.time()
        
        # NOTE: For a real app, we might check if DB exists and add to it.
        # But for this assignment, re-creating it ensures we only key off the CURRENT uploaded files.
        # Let's delete the old one if it exists to be safe.
        if os.path.exists(self.persist_directory):
            try:
                shutil.rmtree(self.persist_directory)
                logger.info("Cleared old database.")
            except Exception as e:
                logger.warning(f"Could not delete old DB: {e}")

        # Create the DB
        vectordb = Chroma.from_documents(
            documents=documents,
            embedding=self.embeddings,
            persist_directory=self.persist_directory
        )
        
        # Force save to disk
        # In newer langchain/chroma versions this handles automatically, but good to be sure
        # vectordb.persist() 
        
        end_time = time.time()
        logger.info(f"Vector DB created in {end_time - start_time:.2f} seconds with {len(documents)} chunks.")
        
        return vectordb
